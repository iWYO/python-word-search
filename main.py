import random
import os
import math
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_CELL_VERTICAL_ALIGNMENT

try:
    from docx2pdf import convert
except ImportError:
    convert = None
    print("WARNING: 'docx2pdf' is not installed. PDF conversion will be skipped.")
    print("Install it using: pip install docx2pdf")

# ==========================================
#              CONFIGURATION
# ==========================================
CONFIG = {
    # --- DIFFICULTY SETTINGS ---
    # Options: "Easy", "Medium", "Hard", "Impossible"
    # The program calculates word count based on this level relative to grid size.
    "DIFFICULTY_LEVEL": "Medium", 

    # --- GRID SETTINGS ---
    "GRID_SIZE": 28,            # Size of the grid (e.g., 28x28)
    
    # --- TECHNICAL ---
    "MAX_ATTEMPTS": 1000,       # Number of attempts to place a word

    # --- FILES & PATHS ---
    "INPUT_FILE": "data/words.txt",  # File containing the word list
    "TEMPLATE_FILE": "data/template.docx", # Word template file
    "OUTPUT_FOLDER": "puzzles", # Folder to save results

    # --- APPEARANCE (GRID) ---
    "FONT_FAMILY": "Courier New", 
    "FONT_SIZE": 14,            # Font size (Pt)
    "ROW_HEIGHT": 14,           # Row height (Pt)
    "COL_WIDTH": 18,            # Column width (Pt)
    
    # --- APPEARANCE (WORD LIST) ---
    "LIST_FONT_SIZE": 9,        # Font size for the word list
    "TITLE_FONT_SIZE": 12,      # Font size for the title
    "LIST_COLUMNS": 4           # Number of columns for the word list
}
# ==========================================

class WordSearchGenerator:
    def __init__(self, size):
        self.size = size
        self.grid = [['' for _ in range(size)] for _ in range(size)]
        self.solutions = []
        # English Alphabet (A-Z)
        self.alphabet = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"

    def place_word(self, word):
        word = word.upper()
        length = len(word)
        directions = [(1, 0), (0, 1), (1, 1), (1, -1)]
        random.shuffle(directions)

        for dx, dy in directions:
            attempts = 0
            while attempts < CONFIG["MAX_ATTEMPTS"]:
                attempts += 1
                
                start_x = random.randint(0, self.size - 1)
                start_y = random.randint(0, self.size - 1)

                end_x = start_x + (length - 1) * dx
                end_y = start_y + (length - 1) * dy

                if not (0 <= end_x < self.size and 0 <= end_y < self.size):
                    continue

                valid_position = True
                for i in range(length):
                    current_x = start_x + i * dx
                    current_y = start_y + i * dy
                    cell = self.grid[current_y][current_x]
                    
                    if cell != '' and cell != word[i]:
                        valid_position = False
                        break
                
                if valid_position:
                    for i in range(length):
                        self.grid[start_y + i * dy][start_x + i * dx] = word[i]
                    self.solutions.append(word)
                    return True
        return False

    def fill_random_chars(self):
        for y in range(self.size):
            for x in range(self.size):
                if self.grid[y][x] == '':
                    self.grid[y][x] = random.choice(self.alphabet)

    def calculate_word_count(self):
        level = CONFIG["DIFFICULTY_LEVEL"]
        size = self.size
        
        base_area = 28 * 28
        current_area = size * size
        scale_factor = current_area / base_area

        if level == "Easy": return 100 # As many as possible
        elif level == "Medium": return max(3, int(20 * scale_factor))
        elif level == "Hard": return max(2, int(10 * scale_factor))
        elif level == "Impossible": return 1 
        else: return 20

    def generate_from_file(self, file_path):
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                all_words = [line.strip() for line in f if line.strip()]
            
            target_count = self.calculate_word_count()
            print(f" -> Target word count for '{CONFIG['DIFFICULTY_LEVEL']}': {target_count}")
            
            if len(all_words) > target_count:
                words = random.sample(all_words, target_count)
            else:
                words = all_words
            
            # Sort by length (descending) to place long words first
            words.sort(key=len, reverse=True)

            placed = []
            skipped = []

            for word in words:
                if self.place_word(word):
                    placed.append(word)
                else:
                    skipped.append(word)
            
            self.fill_random_chars()
            return placed, skipped
        except FileNotFoundError:
            print(f"Error: Input file not found: {file_path}")
            return [], []

    def save_to_docx(self, output_path, template_path=None, puzzle_id=None):
        doc = None
        if template_path and os.path.exists(template_path):
            try:
                doc = Document(template_path)
            except:
                print("Error opening template, creating blank document.")
                doc = Document()
        else:
            doc = Document()

        # --- TAG REPLACEMENT ---
        replacements = {}
        if puzzle_id is not None: replacements["[ID]"] = str(puzzle_id)
        replacements["[DIFFICULTY]"] = CONFIG["DIFFICULTY_LEVEL"]
        
        placed_count = len(self.solutions)
        info_text = f"({self.size}x{self.size} / {placed_count} words)"
        replacements["[INFO]"] = info_text

        def replace_in_all(tag, text):
            def replace_in_paragraph(paragraph):
                if tag in paragraph.text:
                    replaced = False
                    for run in paragraph.runs:
                        if tag in run.text:
                            run.text = run.text.replace(tag, text)
                            replaced = True
                    if not replaced:
                        paragraph.text = paragraph.text.replace(tag, text)

            for p in doc.paragraphs: replace_in_paragraph(p)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs: replace_in_paragraph(p)
            for section in doc.sections:
                for hf in [section.header, section.footer, section.first_page_header, section.first_page_footer]:
                    if hf:
                        for p in hf.paragraphs: replace_in_paragraph(p)

        for tag, text in replacements.items():
            replace_in_all(tag, text)

        # --- 1. THE GRID ---
        placeholder_grid = None
        for p in doc.paragraphs:
            if '[GRID]' in p.text:
                placeholder_grid = p
                break
        
        table = doc.add_table(rows=self.size, cols=self.size)

        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.autofit = False 
        
        row_height = Pt(CONFIG["ROW_HEIGHT"])
        col_width = Pt(CONFIG["COL_WIDTH"])

        for y in range(self.size):
            row = table.rows[y]
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row.height = row_height

            for x in range(self.size):
                cell = row.cells[x]
                cell.width = col_width
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                
                paragraph = cell.paragraphs[0]
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0

                run = paragraph.add_run(self.grid[y][x])
                run.font.size = Pt(CONFIG["FONT_SIZE"])
                run.font.name = CONFIG["FONT_FAMILY"]
                run.font.bold = True
                
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if placeholder_grid:
            placeholder_grid._p.addnext(table._tbl)
            p_element = placeholder_grid._element
            p_element.getparent().remove(p_element)
        
        # --- 2. THE WORD LIST ---
        placeholder_words = None
        for p in doc.paragraphs:
            if '[WORDS]' in p.text:
                placeholder_words = p
                break
        
        # Title
        p_title = doc.add_paragraph()
        run_title = p_title.add_run("Words to find:")
        run_title.bold = True
        run_title.font.size = Pt(CONFIG["TITLE_FONT_SIZE"])
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Table List
        sorted_words = sorted(self.solutions)
        word_count = len(sorted_words)
        num_cols = CONFIG["LIST_COLUMNS"]
        num_rows = math.ceil(word_count / num_cols)

        list_table = doc.add_table(rows=num_rows, cols=num_cols)
        list_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        list_table.autofit = True 

        for index, word in enumerate(sorted_words):
            r = index // num_cols
            c = index % num_cols
            cell = list_table.rows[r].cells[c]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            
            run_box = p.add_run("☐ ")
            run_box.font.size = Pt(CONFIG["LIST_FONT_SIZE"] + 2)
            
            run_word = p.add_run(word)
            run_word.font.size = Pt(CONFIG["LIST_FONT_SIZE"])

        # Insert into Doc
        if placeholder_words:
            placeholder_words._p.addnext(p_title._p)
            p_title._p.addnext(list_table._tbl)
            
            p_element = placeholder_words._element
            p_element.getparent().remove(p_element)
        
        doc.save(output_path)
        print(f"Word document saved: {output_path}")

# --- MAIN PROGRAM ---

if __name__ == "__main__":
    script_dir = os.path.dirname(os.path.abspath(__file__))
    input_file = os.path.join(script_dir, CONFIG["INPUT_FILE"])
    template_file = os.path.join(script_dir, CONFIG["TEMPLATE_FILE"])
    output_folder = os.path.join(script_dir, CONFIG["OUTPUT_FOLDER"])

    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
        print(f"Folder created: {output_folder}")

    max_id = 0
    for f in os.listdir(output_folder):
        if f.startswith("puzzle-") and f.endswith(".docx"):
            try:
                num = int(f.replace("puzzle-", "").replace(".docx", ""))
                if num > max_id: max_id = num
            except ValueError: continue

    next_id = max_id + 1
    docx_name = f"puzzle-{next_id}.docx"
    pdf_name = f"puzzle-{next_id}.pdf"
    
    docx_path = os.path.join(output_folder, docx_name)
    pdf_path = os.path.join(output_folder, pdf_name)

    print(f"--- GENERATION STARTED ---")
    print(f"ID: {next_id} | Difficulty: {CONFIG['DIFFICULTY_LEVEL']}")
    
    generator = WordSearchGenerator(size=CONFIG["GRID_SIZE"]) 
    placed_words, skipped_words = generator.generate_from_file(input_file)

    print(f"Words placed: {len(placed_words)}")
    if skipped_words:
        print(f"Could not place ({len(skipped_words)}): {skipped_words}")

    # 1. Generate Word
    generator.save_to_docx(docx_path, template_path=template_file, puzzle_id=next_id)
    
    # 2. Convert to PDF
    if convert:
        print("Converting to PDF... (this may take a moment)")
        try:
            convert(os.path.abspath(docx_path), os.path.abspath(pdf_path))
            print(f"Success! PDF saved: {pdf_path}")
        except Exception as e:
            print(f"PDF conversion error: {e}")
            print("Ensure Word is not open or stuck.")
    else:
        print("PDF skipped.")